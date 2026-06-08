from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_ROOT = ROOT.parent / "v0.1_技术交付文档"


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    doc.add_paragraph("")


def update_api_doc() -> None:
    path = DOC_ROOT / "方案C_v0.1_API接口文档.docx"
    doc = Document(path)
    style_document(doc)

    doc.add_page_break()
    add_heading(doc, "v1.3 商家入驻与打款凭证接口变更", 1)
    doc.add_paragraph("本章节同步 PRD v1.3 的权限和接口调整：商家端只保留入驻申请、入驻状态、打款凭证列表、打款凭证详情、我的商家信息；客户申请处理由业务员、审核员、出纳和超级管理员完成。")

    add_heading(doc, "新增商家端接口", 2)
    add_table(
        doc,
        ["接口", "方法", "权限", "说明"],
        [
            ["/api/v1/merchant/onboarding", "POST", "STORE", "提交商家入驻申请"],
            ["/api/v1/merchant/me", "GET", "STORE", "获取我的商家信息和最新入驻状态"],
            ["/api/v1/merchant/vouchers", "GET", "STORE", "获取本商家打款凭证列表"],
            ["/api/v1/merchant/vouchers/{voucherId}", "GET", "STORE", "获取本商家打款凭证详情"],
        ],
    )

    add_heading(doc, "新增后台接口", 2)
    add_table(
        doc,
        ["接口", "方法", "权限", "说明"],
        [
            ["/api/v1/admin/merchants", "GET", "SUPER_ADMIN", "获取商家入驻申请列表"],
            ["/api/v1/admin/merchants/{onboardingId}", "GET", "SUPER_ADMIN", "获取商家入驻申请详情"],
            ["/api/v1/admin/merchants/{onboardingId}/approve", "POST", "SUPER_ADMIN", "审核通过商家入驻"],
            ["/api/v1/admin/merchants/{onboardingId}/reject", "POST", "SUPER_ADMIN", "驳回商家入驻"],
            ["/api/v1/admin/merchant-vouchers", "GET", "SUPER_ADMIN", "获取后台打款凭证列表"],
            ["/api/v1/admin/merchant-vouchers", "POST", "SUPER_ADMIN", "新增商家打款凭证"],
            ["/api/v1/admin/merchant-vouchers/{voucherId}/void", "POST", "SUPER_ADMIN", "作废商家打款凭证"],
        ],
    )

    add_heading(doc, "旧接口权限调整", 2)
    add_table(
        doc,
        ["接口", "调整"],
        [
            ["/api/v1/applications", "POST 仅 SUPER_ADMIN 可创建测试申请，商家不可创建客户申请"],
            ["/api/v1/applications", "GET 商家不可访问客户申请列表"],
            ["/api/v1/applications/{id}", "GET 商家不可访问客户申请详情"],
            ["/api/v1/applications/{id}/logs", "GET 商家不可访问客户申请状态日志"],
            ["/api/v1/applications/{id}/supplement", "POST 仅 SALES 和 SUPER_ADMIN 可提交客户申请补资料"],
            ["/api/v1/applications/{id}/request-supplement", "POST 补资料处理对象仅支持 SALES"],
            ["/api/v1/attachments", "POST 商家不可上传客户申请附件"],
        ],
    )

    add_heading(doc, "请求示例：提交商家入驻申请", 2)
    doc.add_paragraph('POST /api/v1/merchant/onboarding')
    doc.add_paragraph('{"merchantName":"东区旗舰店","contactName":"张三","paymentMethod":"BANK","paymentAccount":"6222020202020202020","qualificationFile":{"fileName":"license.pdf","filePath":"/demo/merchant/license.pdf","mimeType":"application/pdf"}}')

    add_heading(doc, "响应示例：商家打款凭证详情", 2)
    doc.add_paragraph('{"success":true,"data":{"voucher":{"voucherNo":"PV20260608120000123","relatedBusinessNo":"A20260531032113666","amount":9000,"status":"PAID","voucherFile":{"fileName":"payout-voucher.png","filePath":"/demo/payout-voucher.png","mimeType":"image/png"}}}}')

    doc.save(path)


def update_schema_doc() -> None:
    path = DOC_ROOT / "方案C_v0.1_数据表设计文档.docx"
    doc = Document(path)
    style_document(doc)

    doc.add_page_break()
    add_heading(doc, "v1.3 商家入驻与打款凭证数据表变更", 1)
    doc.add_paragraph("本章节同步 PRD v1.3 的数据模型调整：新增商家入驻申请表、商家打款凭证表，并在 stores 表补充入驻状态和收款信息。")

    add_heading(doc, "stores 表新增字段", 2)
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["onboarding_status", "varchar(40)", "商家入驻状态：PENDING_REVIEW、APPROVED、REJECTED"],
            ["payment_method", "varchar(40)", "收款方式"],
            ["payment_account", "varchar(120)", "原始收款账号，前端返回需脱敏"],
            ["payment_account_name", "varchar(120)", "收款户名"],
            ["payment_bank_or_channel", "varchar(120)", "开户行或收款渠道"],
        ],
    )

    add_heading(doc, "merchant_onboarding_applications 表", 2)
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["id", "uuid", "主键"],
            ["store_id", "uuid", "关联 stores.id，可为空"],
            ["applicant_name / applicant_phone / applicant_id_number", "varchar", "申请人信息"],
            ["merchant_name / merchant_address", "varchar", "商家名称和地址"],
            ["contact_name / contact_phone", "varchar", "联系人信息"],
            ["payment_method / payment_account / payment_account_name / payment_bank_or_channel", "varchar", "收款信息"],
            ["id_card_front_file / id_card_back_file / qualification_file", "json", "入驻资料文件元数据"],
            ["status", "varchar(40)", "入驻审核状态"],
            ["reviewer_user_id / reviewed_at / review_note / reject_reason", "mixed", "审核记录"],
        ],
    )

    add_heading(doc, "merchant_payment_vouchers 表", 2)
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["id", "uuid", "主键"],
            ["voucher_no", "varchar(40)", "凭证编号，唯一"],
            ["store_id", "uuid", "关联 stores.id"],
            ["payout_record_id", "uuid", "关联 payout_records.id，可为空"],
            ["related_business_no", "varchar(80)", "关联业务编号"],
            ["amount", "decimal(12,2)", "打款金额"],
            ["status", "varchar(40)", "PENDING_CONFIRMATION、PAID、VOIDED"],
            ["paid_at", "timestamp", "打款时间"],
            ["payee_name / payee_account_masked / payer_name", "varchar", "收款方和付款方信息"],
            ["voucher_file", "json", "凭证文件元数据"],
            ["remark / void_reason", "text", "备注和作废原因"],
            ["created_by_user_id", "bigint", "创建人"],
        ],
    )

    add_heading(doc, "旧表业务口径调整", 2)
    add_table(
        doc,
        ["表", "调整"],
        [
            ["applications", "商家不再创建或查看客户完整申请，测试申请由超级管理员或种子数据生成"],
            ["status_logs", "客户申请补资料处理对象不再写入商家，改为业务员"],
            ["attachments", "商家不再上传客户申请附件，商家资料文件以 JSON 元数据保存到商家表"],
            ["payout_records", "出纳确认打款后保留原记录，并生成 merchant_payment_vouchers 供商家查看"],
        ],
    )

    doc.save(path)


if __name__ == "__main__":
    update_api_doc()
    update_schema_doc()
