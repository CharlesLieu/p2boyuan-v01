from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "v0.1_UAT测试启动说明.md"
OUTPUT = ROOT.parent / "v0.1_技术交付文档" / "v0.1_UAT测试启动说明.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(9)
    run.bold = bold


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for index, text in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, text, bold=True)
        set_cell_shading(cell, "EAF4F1")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows[1:]:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            set_cell_text(cells[index], text)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph("")


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(10.5)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(10)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        if not set(raw.replace("|", "").replace("-", "").replace(":", "").strip()):
            i += 1
            continue
        cells = [cell.strip().replace("`", "") for cell in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(12)
            run = paragraph.add_run(line[2:])
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(33, 45, 58)
        elif line.startswith("## "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(line[3:])
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 110, 92)
        elif line.startswith("### "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(line[4:])
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(33, 45, 58)
        elif line.startswith("- "):
            add_bullet(doc, line[2:])
        elif line.startswith("|"):
            table_rows, next_i = parse_markdown_table(lines, i)
            add_table(doc, table_rows)
            i = next_i
            continue
        else:
            add_paragraph(doc, line.replace("`", ""))

        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
